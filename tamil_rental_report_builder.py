import json
import os
import re
from datetime import date
from html import escape

from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    KeepTogether,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = os.path.abspath(os.path.dirname(__file__))
SOURCE_JSON = os.path.join(ROOT, "output", "pdf", "rathinam_rental_research_summary.json")
OUTDIR = os.path.join(ROOT, "output", "tamil_report")
ASSET_DIR = os.path.join(ROOT, "assets")
os.makedirs(OUTDIR, exist_ok=True)

REGULAR_FONT = os.path.join(ASSET_DIR, "NotoSansTamil-Regular.ttf")
BOLD_FONT = os.path.join(ASSET_DIR, "NotoSansTamil-Bold.ttf")
LATIN_FONT = os.path.join(ASSET_DIR, "NotoSans-Regular.ttf")
LATIN_BOLD_FONT = os.path.join(ASSET_DIR, "NotoSans-Bold.ttf")

TODAY_TAMIL = "05 à®œà¯‚à®²à¯ˆ 2026"
BLUE = colors.HexColor("#123A63")
MID_BLUE = colors.HexColor("#2E74B5")
LIGHT_BLUE = colors.HexColor("#EAF2FB")
PALE_BLUE = colors.HexColor("#F5F9FE")
INK = colors.HexColor("#111827")
MUTED = colors.HexColor("#5B677A")
GRID = colors.HexColor("#C9D3E1")


def load_data():
    with open(SOURCE_JSON, "r", encoding="utf-8") as f:
        data = json.load(f)
    return data["ranked"]


def rupee(value):
    if isinstance(value, int):
        return f"â‚¹{value:,}"
    if isinstance(value, str):
        return value.replace("Rs ", "â‚¹")
    return str(value)


AREA_SCORES = [
    ["Eachanari", 8, 10, 8, 8, 8, 10, 8, 8],
    ["Malumichampatti", 7, 8, 7, 7, 7, 8, 9, 8],
    ["Kurichi", 8, 7, 9, 8, 8, 7, 7, 7],
    ["Podanur", 7, 6, 9, 9, 8, 6, 7, 6],
    ["Kuniyamuthur", 8, 6, 9, 8, 9, 6, 6, 7],
]


def area_table_rows():
    rows = [["à®ªà®•à¯à®¤à®¿", "à®ªà®¾à®¤à¯à®•à®¾à®ªà¯à®ªà¯", "à®¤à¯‚à®°à®®à¯", "à®‰à®£à®µà¯", "à®®à®°à¯à®¤à¯à®¤à¯à®µà®®à¯", "Internet", "à®ªà®¯à®£à®®à¯", "à®µà®¾à®Ÿà®•à¯ˆ", "à®…à®®à¯ˆà®¤à®¿", "à®®à¯Šà®¤à¯à®¤à®®à¯"]]
    for area, safety, distance, food, hospital, internet, commute, rent, peace in AREA_SCORES:
        total = round((safety + distance + food + hospital + internet + commute + rent + peace) / 8, 1)
        rows.append([area, safety, distance, food, hospital, internet, commute, rent, peace, total])
    return rows


def house_advantages(row, idx):
    if idx == 0:
        return "à®•à¯à®±à¯ˆà®¨à¯à®¤ à®µà®¾à®Ÿà®•à¯ˆ, à®®à®¿à®- à®…à®°à¯à®•à®¿à®²à¯, Semi-Furnished, à®•à¯à®±à¯ˆà®¨à¯à®¤ Security Deposit"
    if row["locality"] in ("Eachanari",):
        return "à®•à®²à¯à®²à¯‚à®°à®¿à®•à¯à®•à¯ à®®à®¿à®- à®…à®°à¯à®•à®¿à®²à¯, commute à®Žà®³à®¿à®¤à¯"
    if row["locality"] in ("Sugunapuram", "Madukkarai"):
        return "à®µà®¾à®Ÿà®•à¯ˆ à®•à®Ÿà¯à®Ÿà¯à®ªà¯à®ªà®¾à®Ÿà¯à®Ÿà®¿à®²à¯, à®…à®®à¯ˆà®¤à®¿à®¯à®¾à®© residential pocket à®†à®- à®‡à®°à¯à®•à¯à®- à®µà®¾à®¯à¯à®ªà¯à®ªà¯"
    if row["locality"] == "Sundarapuram":
        return "à®•à®Ÿà¯ˆà®•à®³à¯, à®‰à®£à®µà¯, services à®Žà®³à®¿à®¤à®¿à®²à¯ à®•à®¿à®Ÿà¯ˆà®•à¯à®•à¯à®®à¯"
    return "à®¨à®•à®° à®µà®šà®¤à®¿à®•à®³à¯ à®¨à®©à¯à®±à®¾à®- à®•à®¿à®Ÿà¯ˆà®•à¯à®•à¯à®®à¯"


def house_disadvantages(row, idx):
    deposit = row["deposit"]
    if idx == 0:
        return "Water, parking, owner terms à®¨à¯‡à®°à®¿à®²à¯ à®‰à®±à¯à®¤à®¿ à®šà¯†à®¯à¯à®¯ à®µà¯‡à®£à¯à®Ÿà¯à®®à¯"
    if deposit == "Not public":
        return "Security Deposit public à®†à®- à®‡à®²à¯à®²à¯ˆ; owner-à® à®•à¯‡à®Ÿà¯à®Ÿà¯‡ à®‰à®±à¯à®¤à®¿ à®šà¯†à®¯à¯à®¯ à®µà¯‡à®£à¯à®Ÿà¯à®®à¯"
    if "500,000" in deposit:
        return "Deposit à®®à®¿à®- à®…à®¤à®¿à®•à®®à®¾à®- à®¤à¯†à®°à®¿à®•à®¿à®±à®¤à¯; avoid à®…à®²à¯à®²à®¤à¯ à®•à®Ÿà¯à®®à¯ˆà®¯à®¾à®- verify à®šà¯†à®¯à¯à®¯ à®µà¯‡à®£à¯à®Ÿà¯à®®à¯"
    if row["furnished"] == "Unfurnished":
        return "Unfurnished; mattress/utensils setup à®šà¯†à®²à®µà¯ à®µà®°à¯à®®à¯"
    return "Final street safety à®®à®±à¯à®±à¯à®®à¯ water supply verify à®šà¯†à®¯à¯à®¯ à®µà¯‡à®£à¯à®Ÿà¯à®®à¯"


def house_table_rows(houses):
    rows = [["#", "à®µà¯€à®Ÿà¯", "à®µà®¾à®Ÿà®•à¯ˆ", "Security Deposit", "à®¤à¯‚à®°à®®à¯", "à®ªà®¯à®£ à®¨à¯‡à®°à®®à¯", "Furnished", "Parking", "à®¨à®©à¯à®®à¯ˆà®•à®³à¯", "à®•à®µà®©à®¿à®•à¯à®- à®µà¯‡à®£à¯à®Ÿà®¿à®¯à®¤à¯"]]
    for i, h in enumerate(houses, 1):
        short_type = "1BHK House" if h["bedrooms"] == "1" and "House" in h["type"] else (
            "2BHK House" if h["bedrooms"] == "2" and "House" in h["type"] else f"{h['bedrooms']}BHK Flat"
        )
        rows.append([
            str(i),
            f"{h['locality']} - {short_type}",
            rupee(h["rent"]),
            rupee(h["deposit"]),
            f"{h['distance']} km",
            h["time"],
            "Semi-Furnished" if h["furnished"] == "Semi-Furnished" else "Unfurnished",
            h["parking"].replace("Ask owner", "Owner-à® à®•à¯‡à®Ÿà¯à®- à®µà¯‡à®£à¯à®Ÿà¯à®®à¯"),
            house_advantages(h, i - 1),
            house_disadvantages(h, i - 1),
        ])
    return rows


MONTHLY_BUDGET = [
    ["à®µà®¾à®Ÿà®•à¯ˆ", 7000, 6500, 10000],
    ["à®®à®¿à®©à¯à®šà®¾à®°à®®à¯ (Electricity)", 1000, 600, 1500],
    ["à®‡à®£à¯ˆà®¯à®®à¯ (Fiber Internet / Wi-Fi)", 900, 600, 1200],
    ["à®®à®³à®¿à®•à¯ˆ + à®šà®®à¯ˆà®ªà¯à®ªà¯", 5000, 3500, 6500],
    ["Bike fuel + maintenance", 1800, 1200, 2500],
    ["Emergency / medical buffer", 1500, 800, 2500],
]

ONE_TIME = [
    ["Security Deposit", "â‚¹10,000", "Recommended house à®…à®Ÿà®¿à®ªà¯à®ªà®Ÿà¯ˆà®¯à®¿à®²à¯"],
    ["Kitchen essentials", "â‚¹3,500 - â‚¹6,000", "Induction/gas, à®ªà®¾à®¤à¯à®¤à®¿à®°à®™à¯à®•à®³à¯, storage"],
    ["Mattress", "â‚¹3,000 - â‚¹6,000", "à®µà¯€à®Ÿà¯à®Ÿà®¿à®²à¯ bed à®‡à®²à¯à®²à¯ˆà®¯à¯†à®©à®¿à®²à¯"],
    ["Wi-Fi setup/router", "â‚¹0 - â‚¹2,000", "ISP offer à®ªà¯Šà®±à¯à®¤à¯à®¤à¯"],
    ["Basic utensils", "â‚¹2,000 - â‚¹3,500", "Plate, tumbler, cooker, pan"],
    ["Cleaning supplies", "â‚¹800 - â‚¹1,500", "Bucket, broom, mop, liquids"],
    ["Bike accessories", "â‚¹1,000 - â‚¹2,000", "Lock, rain cover, basic toolkit"],
]

VERIFY_ITEMS = [
    "à®¤à®£à¯à®£à¯€à®°à¯ source - borewell / corporation / tanker backup",
    "Fiber Internet à®•à®¿à®Ÿà¯ˆà®•à¯à®•à®¿à®±à®¤à®¾ - Airtel / Jio / ACT / local ISP",
    "à®…à®°à¯à®•à®¿à®²à¯ à®‰à®³à¯à®³ neighbours à®Žà®ªà¯à®ªà®Ÿà®¿ à®‡à®°à¯à®•à¯à®•à®¿à®±à®¾à®°à¯à®•à®³à¯",
    "Bike parking à®ªà®¾à®¤à¯à®•à®¾à®ªà¯à®ªà®¾à®©à®¤à®¾",
    "Rent agreement, notice period, lock-in period",
    "Security Deposit refund terms à®Žà®´à¯à®¤à¯à®¤à®¿à®²à¯ à®‰à®³à¯à®³à®¤à®¾",
    "Power cuts à®Žà®µà¯à®µà®³à®µà¯ à®µà®°à¯à®•à®¿à®±à®¤à¯",
    "à®‡à®°à®µà¯ street safety à®®à®±à¯à®±à¯à®®à¯ street light",
    "Noise - traffic, workshop, factory, loudspeaker à®ªà¯‹à®©à¯à®±à®µà¯ˆ",
    "Mobile signal indoor speed test",
    "Owner documents à®®à®±à¯à®±à¯à®®à¯ ID proof",
    "EB connection à®¤à®©à®¿à®¯à®¾à®- à®‰à®³à¯à®³à®¤à®¾; meter reading à®Žà®ªà¯à®ªà®Ÿà®¿ à®•à®£à®•à¯à®•à®¿à®Ÿà®ªà¯à®ªà®Ÿà¯à®®à¯",
]

RISKS = [
    ["à®¤à®©à®¿à®¯à®¾à®- à®µà®¾à®´à¯à®µà®¤à¯", "Daily routine, emergency contacts, location sharing à®µà¯ˆà®¤à¯à®¤à¯à®•à¯à®•à¯Šà®³à¯à®µà®¤à¯."],
    ["à®šà®®à¯ˆà®ªà¯à®ªà®¤à¯", "Simple cooking plan: rice, eggs, vegetables, curd, tea. à®µà¯†à®³à®¿ à®‰à®£à®µà¯ˆ à®•à¯à®±à¯ˆà®ªà¯à®ªà®¤à¯."],
    ["Bills", "Rent, EB, Wi-Fi due dates-à® calendar reminder-à®²à¯ à®ªà¯‹à®Ÿà¯à®µà®¤à¯."],
    ["Cleaning", "à®µà®¾à®°à®®à¯ à®‡à®°à¯à®®à¯à®±à¯ˆ room cleaning; à®®à®¾à®¤à®®à¯ à®’à®°à¯à®®à¯à®±à¯ˆ deep cleaning."],
    ["Maintenance", "Owner-à®•à¯à®•à¯ WhatsApp proof à®‰à®Ÿà®©à¯ à®‰à®Ÿà®©à¯‡ à®¤à¯†à®°à®¿à®µà®¿à®¤à¯à®¤à¯ repair record à®µà¯ˆà®¤à¯à®¤à¯à®•à¯à®•à¯Šà®³à¯à®µà®¤à¯."],
    ["Emergency", "Nearby hospital, pharmacy, à®¨à®£à¯à®ªà®°à¯, owner contact numbers phone-à®²à¯ pin à®šà¯†à®¯à¯à®µà®¤à¯."],
]


def tamil_content(houses):
    best = houses[0]
    expected_monthly = sum(row[1] for row in MONTHLY_BUDGET)
    min_monthly = sum(row[2] for row in MONTHLY_BUDGET)
    max_monthly = sum(row[3] for row in MONTHLY_BUDGET)
    yearly = expected_monthly * 12 + 17000
    return {
        "best": best,
        "expected_monthly": expected_monthly,
        "min_monthly": min_monthly,
        "max_monthly": max_monthly,
        "yearly": yearly,
        "intro": (
            "à®‡à®¨à¯à®¤ à®…à®±à®¿à®•à¯à®•à¯ˆ à®°à®¤à¯à®¤à®¿à®©à®®à¯ Technical Campus à®…à®°à¯à®•à®¿à®²à¯ à®šà¯à®®à®¾à®°à¯ à®’à®°à¯ à®µà®°à¯à®Ÿà®¤à¯à®¤à®¿à®±à¯à®•à¯ à®®à®Ÿà¯à®Ÿà¯à®®à¯ "
            "à®µà®¾à®Ÿà®•à¯ˆ à®µà¯€à®Ÿà¯à®Ÿà¯ˆ à®¤à¯‡à®°à¯à®µà¯ à®šà¯†à®¯à¯à®µà®¤à®±à¯à®•à®¾à®© à®¤à¯†à®³à®¿à®µà®¾à®© à®†à®¯à¯à®µà¯. à®¨à¯‹à®•à¯à®•à®®à¯ à®µà¯†à®±à¯à®®à¯ à®µà®šà®¤à®¿ à®…à®²à¯à®²; "
            "à®ªà®Ÿà®¿à®ªà¯à®ªà¯, AI development, software projects, MBA preparation à®†à®•à®¿à®¯à®µà®±à¯à®±à®¿à®±à¯à®•à¯ à®…à®®à¯ˆà®¤à®¿à®¯à®¾à®© "
            "à®®à®±à¯à®±à¯à®®à¯ à®•à®Ÿà¯à®Ÿà¯à®ªà¯à®ªà®¾à®Ÿà¯à®Ÿà®¾à®© à®šà¯‚à®´à®²à¯ˆ à®‰à®°à¯à®µà®¾à®•à¯à®•à¯à®µà®¤à¯."
        ),
    }


def pdf_styles():
    pdfmetrics.registerFont(TTFont("NotoTamil", REGULAR_FONT))
    pdfmetrics.registerFont(TTFont("NotoTamil-Bold", BOLD_FONT))
    pdfmetrics.registerFont(TTFont("NotoLatin", LATIN_FONT))
    pdfmetrics.registerFont(TTFont("NotoLatin-Bold", LATIN_BOLD_FONT))
    styles = getSampleStyleSheet()
    style_map = {
        "title": ParagraphStyle("TamilTitle", parent=styles["Title"], fontName="NotoTamil-Bold", fontSize=22, leading=30, textColor=BLUE, alignment=TA_CENTER, spaceAfter=12),
        "subtitle": ParagraphStyle("TamilSubtitle", parent=styles["BodyText"], fontName="NotoTamil", fontSize=11, leading=17, textColor=MUTED, alignment=TA_CENTER, spaceAfter=10),
        "h1": ParagraphStyle("TamilH1", parent=styles["Heading1"], fontName="NotoTamil-Bold", fontSize=15, leading=21, textColor=BLUE, spaceBefore=8, spaceAfter=7),
        "h2": ParagraphStyle("TamilH2", parent=styles["Heading2"], fontName="NotoTamil-Bold", fontSize=12, leading=18, textColor=MID_BLUE, spaceBefore=6, spaceAfter=5),
        "body": ParagraphStyle("TamilBody", parent=styles["BodyText"], fontName="NotoTamil", fontSize=9.6, leading=15, textColor=INK, spaceAfter=5),
        "small": ParagraphStyle("TamilSmall", parent=styles["BodyText"], fontName="NotoTamil", fontSize=8, leading=12, textColor=INK),
        "tiny": ParagraphStyle("TamilTiny", parent=styles["BodyText"], fontName="NotoTamil", fontSize=6.8, leading=9, textColor=INK),
        "cover_meta": ParagraphStyle("CoverMeta", parent=styles["BodyText"], fontName="NotoTamil", fontSize=10, leading=16, textColor=INK, alignment=TA_CENTER),
        "card": ParagraphStyle("Card", parent=styles["BodyText"], fontName="NotoTamil", fontSize=10, leading=16, textColor=INK, spaceAfter=4),
    }
    style_map["small_white"] = ParagraphStyle("TamilSmallWhite", parent=style_map["small"], textColor=colors.white)
    style_map["tiny_white"] = ParagraphStyle("TamilTinyWhite", parent=style_map["tiny"], textColor=colors.white)
    return style_map


def P(text, style):
    raw = escape(str(text))
    def repl(match):
        chunk = match.group(0)
        return f'<font name="NotoLatin">{chunk}</font>'
    mixed = re.sub(r"[A-Za-z0-9â‚¹,.;:/()|&=\-+% ]+", repl, raw)
    return Paragraph(mixed, style)


def make_table(rows, col_widths, style, header=True, tiny=False):
    cell_style = style["tiny"] if tiny else style["small"]
    header_style = style["tiny_white"] if tiny else style["small_white"]
    wrapped = []
    for r_idx, row in enumerate(rows):
        wrapped.append([P(c, header_style if header and r_idx == 0 else cell_style) for c in row])
    table = Table(wrapped, colWidths=col_widths, repeatRows=1 if header else 0)
    commands = [
        ("GRID", (0, 0), (-1, -1), 0.35, GRID),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, PALE_BLUE]),
    ]
    if header:
        commands += [
            ("BACKGROUND", (0, 0), (-1, 0), BLUE),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "NotoTamil-Bold"),
        ]
    table.setStyle(TableStyle(commands))
    return table


def add_card(story, title, lines, style):
    rows = [[P(title, style["h2"])]]
    for line in lines:
        rows.append([P(line, style["card"])])
    table = Table(rows, colWidths=[176 * mm])
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), LIGHT_BLUE),
        ("BOX", (0, 0), (-1, -1), 1.0, MID_BLUE),
        ("LEFTPADDING", (0, 0), (-1, -1), 9),
        ("RIGHTPADDING", (0, 0), (-1, -1), 9),
        ("TOPPADDING", (0, 0), (-1, -1), 7),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
    ]))
    story.append(table)
    story.append(Spacer(1, 5 * mm))


def build_pdf(houses):
    style = pdf_styles()
    info = tamil_content(houses)
    path = os.path.join(OUTDIR, "rathinam_house_report_tamil.pdf")
    doc = SimpleDocTemplate(path, pagesize=A4, leftMargin=16 * mm, rightMargin=16 * mm, topMargin=15 * mm, bottomMargin=14 * mm)
    story = []

    story.append(Spacer(1, 28 * mm))
    story.append(P("à®°à®¤à¯à®¤à®¿à®©à®®à¯ à®•à®²à¯à®²à¯‚à®°à®¿ à®…à®°à¯à®•à®¿à®²à¯à®³à¯à®³ à®µà®¾à®Ÿà®•à¯ˆ à®µà¯€à®Ÿà¯", style["title"]))
    story.append(P("à®®à¯à®´à¯à®®à¯ˆà®¯à®¾à®© à®†à®¯à¯à®µà¯ à®…à®±à®¿à®•à¯à®•à¯ˆ", style["title"]))
    story.append(Spacer(1, 10 * mm))
    story.append(P("Prepared for: My Father", style["cover_meta"]))
    story.append(P("Prepared by: Dragon", style["cover_meta"]))
    story.append(P(f"Date: {TODAY_TAMIL}", style["cover_meta"]))
    story.append(Spacer(1, 18 * mm))
    add_card(story, "à®®à¯à®•à¯à®•à®¿à®¯ à®®à¯à®Ÿà®¿à®µà¯", [
        "à®šà¯à®®à®¾à®°à¯ à®’à®°à¯ à®µà®°à¯à®Ÿà®¤à¯à®¤à®¿à®±à¯à®•à¯ à®®à®Ÿà¯à®Ÿà¯à®®à¯, à®•à¯à®±à¯ˆà®¨à¯à®¤ à®šà¯†à®²à®µà¯ + à®…à®®à¯ˆà®¤à®¿à®¯à®¾à®© à®šà¯‚à®´à®²à¯ + à®µà®¿à®°à¯ˆà®µà®¾à®© commute à®Žà®©à¯à®± à®…à®Ÿà®¿à®ªà¯à®ªà®Ÿà¯ˆà®¯à®¿à®²à¯ Eachanari à®ªà®•à¯à®¤à®¿à®¯à®¿à®²à¯ à®‰à®³à¯à®³ 1 BHK Residential House à®šà®¿à®±à®¨à¯à®¤ à®¤à¯‡à®°à¯à®µà®¾à®- à®¤à¯†à®°à®¿à®•à®¿à®±à®¤à¯.",
        f"à®µà®¾à®Ÿà®•à¯ˆ: {rupee(info['best']['rent'])} | Security Deposit: {rupee(info['best']['deposit'])} | à®ªà®¯à®£à®®à¯: {info['best']['distance']} km / {info['best']['time']}",
    ], style)
    story.append(P("à®‡à®¨à¯à®¤ à®…à®±à®¿à®•à¯à®•à¯ˆ à®‰à®£à®°à¯à®šà¯à®šà®¿à®µà®šà®ªà¯à®ªà®Ÿà¯à®Ÿ à®®à¯à®Ÿà®¿à®µà¯ à®…à®²à¯à®². Public rental portals-à®²à¯ à®•à®¿à®Ÿà¯ˆà®¤à¯à®¤ à®¤à®•à®µà®²à¯à®•à®³à¯, budget, commute, à®ªà®¾à®¤à¯à®•à®¾à®ªà¯à®ªà¯, internet, future MBA plan à®†à®•à®¿à®¯à®µà®±à¯à®±à¯ˆ à®µà¯ˆà®¤à¯à®¤à¯ practical à®†à®¯à¯à®µà®¾à®- à®¤à®¯à®¾à®°à®¿à®•à¯à®•à®ªà¯à®ªà®Ÿà¯à®Ÿà®¤à¯.", style["body"]))
    story.append(PageBreak())

    story.append(P("1. Hostel-à®²à®¿à®°à¯à®¨à¯à®¤à¯ à®µà¯†à®³à®¿à®¯à¯‡à®± à®µà®¿à®°à¯à®®à¯à®ªà¯à®®à¯ à®•à®¾à®°à®£à®®à¯", style["h1"]))
    story.append(P(info["intro"], style["body"]))
    for item in [
        "à®…à®®à¯ˆà®¤à®¿à®¯à®¾à®© à®šà¯‚à®´à®²à®¿à®²à¯ better concentration à®•à®¿à®Ÿà¯ˆà®•à¯à®•à¯à®®à¯.",
        "Hostel noise à®•à¯à®±à¯ˆà®¯à¯à®µà®¤à®¾à®²à¯ à®¨à®²à¯à®² à®¤à¯‚à®•à¯à®•à®®à¯ à®®à®±à¯à®±à¯à®®à¯ à®¨à®¿à®²à¯ˆà®¯à®¾à®© routine à®‰à®°à¯à®µà®¾à®•à¯à®®à¯.",
        "à®‰à®£à®µà¯ control à®®à¯‡à®®à¯à®ªà®Ÿà¯à®®à¯; health à®®à®±à¯à®±à¯à®®à¯ energy à®¨à®¿à®²à¯ˆà®¯à®¾à®- à®‡à®°à¯à®•à¯à®•à¯à®®à¯.",
        "AI development, software projects, portfolio building à®†à®•à®¿à®¯à®µà®±à¯à®±à®¿à®±à¯à®•à¯ privacy à®®à®±à¯à®±à¯à®®à¯ uninterrupted time à®•à®¿à®Ÿà¯ˆà®•à¯à®•à¯à®®à¯.",
        "MBA abroad preparation-à®•à¯à®•à¯ reading, applications, interviews à®†à®•à®¿à®¯à®µà®±à¯à®±à®¿à®²à¯ focus à®…à®¤à®¿à®•à®°à®¿à®•à¯à®•à¯à®®à¯.",
        "à®‡à®¤à¯ à®¨à®¿à®°à®¨à¯à®¤à®° à®®à®¾à®±à¯à®±à®®à¯ à®…à®²à¯à®²; à®šà¯à®®à®¾à®°à¯ 10-12 à®®à®¾à®¤à®™à¯à®•à®³à¯à®•à¯à®•à®¾à®© temporary productivity investment.",
    ]:
        story.append(P("â€¢ " + item, style["body"]))

    story.append(P("2. à®Žà®©à¯ à®¤à¯‡à®µà¯ˆà®•à®³à¯", style["h1"]))
    reqs = [
        "â‚¹5,000-â‚¹10,000 à®µà®¾à®Ÿà®•à¯ˆ; à®®à®¿à®- à®¨à®²à¯à®² option à®Žà®©à¯à®±à®¾à®²à¯ à®®à®Ÿà¯à®Ÿà¯à®®à¯‡ â‚¹12,000 à®µà®°à¯ˆ.",
        "Independent room / 1 RK / 1 BHK / à®šà®¿à®±à®¿à®¯ independent house.",
        "PG, Hostel, shared room, dormitory à®µà¯‡à®£à¯à®Ÿà®¾à®®à¯.",
        "Safe area, peaceful neighbourhood, good roads.",
        "Hospital, grocery, pharmacy, tea shop à®ªà¯‹à®©à¯à®± daily facilities à®…à®°à¯à®•à®¿à®²à¯ à®‡à®°à¯à®•à¯à®- à®µà¯‡à®£à¯à®Ÿà¯à®®à¯.",
        "Bike travel 15-20 minutes-à®•à¯à®•à¯ à®•à¯€à®´à¯ à®‡à®°à¯à®•à¯à®- à®µà¯‡à®£à¯à®Ÿà¯à®®à¯.",
        "à®¨à®²à¯à®² à®¤à®£à¯à®£à¯€à®°à¯, à®¨à®²à¯à®² electricity, à®¨à®²à¯à®² mobile signal, Fiber Internet à®µà®¾à®¯à¯à®ªà¯à®ªà¯.",
        "à®’à®°à¯ à®µà®°à¯à®Ÿà®¤à¯à®¤à®¿à®±à¯à®•à¯ à®ªà®¿à®±à®•à¯ easy move-out à®‡à®°à¯à®•à¯à®- à®µà¯‡à®£à¯à®Ÿà¯à®®à¯.",
    ]
    for item in reqs:
        story.append(P("âœ“ " + item, style["body"]))

    story.append(P("3. à®ªà®•à¯à®¤à®¿à®•à®³à¯ à®’à®ªà¯à®ªà¯€à®Ÿà¯", style["h1"]))
    story.append(make_table(area_table_rows(), [30*mm, 16*mm, 14*mm, 14*mm, 18*mm, 17*mm, 16*mm, 16*mm, 17*mm, 18*mm], style, tiny=True))
    story.append(Spacer(1, 4 * mm))
    story.append(P("à®®à®¤à®¿à®ªà¯à®ªà¯€à®Ÿà¯ public information, commute logic, residential suitability, facilities availability à®†à®•à®¿à®¯à®µà®±à¯à®±à®¿à®©à¯ à®…à®Ÿà®¿à®ªà¯à®ªà®Ÿà¯ˆà®¯à®¿à®²à¯ 10-à®²à¯ à®•à®£à®¿à®•à¯à®•à®ªà¯à®ªà®Ÿà¯à®Ÿà®¤à¯. à®¨à¯‡à®°à®¿à®²à¯ à®ªà®¾à®°à¯à®¤à¯à®¤ à®ªà®¿à®±à®•à¯ final confirmation à®…à®µà®šà®¿à®¯à®®à¯.", style["small"]))

    story.append(PageBreak())
    story.append(P("4. Top 10 à®µà¯€à®Ÿà¯à®•à®³à¯", style["h1"]))
    story.append(make_table(house_table_rows(houses), [7*mm, 30*mm, 14*mm, 18*mm, 13*mm, 15*mm, 18*mm, 18*mm, 30*mm, 32*mm], style, tiny=True))

    story.append(PageBreak())
    story.append(P("5. à®ªà®°à®¿à®¨à¯à®¤à¯à®°à¯ˆà®•à¯à®•à®ªà¯à®ªà®Ÿà¯à®®à¯ à®µà¯€à®Ÿà¯", style["h1"]))
    best = info["best"]
    add_card(story, "Recommended House - Eachanari 1 BHK Residential House", [
        f"à®µà®¾à®Ÿà®•à¯ˆ: {rupee(best['rent'])}",
        f"Security Deposit: {rupee(best['deposit'])}",
        f"Bike distance: {best['distance']} km",
        f"Travel time: {best['time']}",
        f"Furnishing: {best['furnished']}",
        "à®•à®²à¯à®²à¯‚à®°à®¿à®•à¯à®•à¯ à®…à®°à¯à®•à®¿à®²à¯ à®‡à®°à¯à®ªà¯à®ªà®¤à®¾à®²à¯ à®¤à®¿à®©à®šà®°à®¿ commute à®šà¯à®²à®ªà®®à¯.",
        "Semi-Furnished à®Žà®©à¯à®ªà®¤à®¾à®²à¯ à®†à®°à®®à¯à®ª setup à®šà¯†à®²à®µà¯ à®•à¯à®±à¯ˆà®¯à¯à®®à¯.",
        "Temporary one-year stay-à®•à¯à®•à¯ deposit à®®à®±à¯à®±à¯à®®à¯ rent à®‡à®°à®£à¯à®Ÿà¯à®®à¯ practical à®†à®- à®‡à®°à¯à®•à¯à®•à®¿à®±à®¤à¯.",
    ], style)
    story.append(P("à®‡à®¨à¯à®¤ à®µà¯€à®Ÿà¯ à®…à®¤à®¿à®- à®µà®šà®¤à®¿à®•à¯à®•à®¾à®- à®…à®²à¯à®²; à®•à¯à®±à¯ˆà®¨à¯à®¤ à®šà¯†à®²à®µà®¿à®²à¯ à®…à®®à¯ˆà®¤à®¿à®¯à®¾à®©, productive à®šà¯‚à®´à®²à¯ˆ à®‰à®°à¯à®µà®¾à®•à¯à®•à¯à®µà®¤à®±à¯à®•à®¾à®- à®¤à¯‡à®°à¯à®µà¯ à®šà¯†à®¯à¯à®¯à®ªà¯à®ªà®Ÿà¯à®•à®¿à®±à®¤à¯. Final advance à®•à¯Šà®Ÿà¯à®•à¯à®•à¯à®®à¯ à®®à¯à®©à¯ water, parking, agreement, deposit refund, neighbours, internet availability à®†à®•à®¿à®¯à®µà¯ˆ à®¨à¯‡à®°à®¿à®²à¯ verify à®šà¯†à®¯à¯à®¯ à®µà¯‡à®£à¯à®Ÿà¯à®®à¯.", style["body"]))

    story.append(P("6. à®Žà®¤à®¿à®°à¯à®ªà®¾à®°à¯à®•à¯à®•à®ªà¯à®ªà®Ÿà¯à®®à¯ à®®à®¾à®¤à®šà¯ à®šà¯†à®²à®µà¯", style["h1"]))
    budget_rows = [["à®šà¯†à®²à®µà¯", "à®¤à®±à¯à®ªà¯‹à®¤à¯ˆà®¯ à®®à®¤à®¿à®ªà¯à®ªà¯", "à®•à¯à®±à¯ˆà®¨à¯à®¤à®ªà®Ÿà¯à®šà®®à¯", "à®…à®¤à®¿à®•à®ªà®Ÿà¯à®šà®®à¯"]]
    for name, current, minimum, maximum in MONTHLY_BUDGET:
        budget_rows.append([name, rupee(current), rupee(minimum), rupee(maximum)])
    budget_rows.append(["à®®à¯Šà®¤à¯à®¤à®®à¯", rupee(info["expected_monthly"]), rupee(info["min_monthly"]), rupee(info["max_monthly"])])
    story.append(make_table(budget_rows, [62*mm, 38*mm, 38*mm, 38*mm], style))
    story.append(P(f"à®’à®°à¯ à®µà®°à¯à®Ÿ à®šà¯†à®²à®µà¯ à®•à®£à®¿à®ªà¯à®ªà¯: à®®à®¾à®¤ à®šà®°à®¾à®šà®°à®¿ {rupee(info['expected_monthly'])} Ã— 12 + basic setup buffer à®šà¯‡à®°à¯à®¤à¯à®¤à¯ à®šà¯à®®à®¾à®°à¯ {rupee(info['yearly'])}.", style["body"]))

    story.append(P("7. One-Time Expenses", style["h1"]))
    story.append(make_table([["Item", "Approx. Amount", "à®•à¯à®±à®¿à®ªà¯à®ªà¯"]] + ONE_TIME + [["à®®à¯Šà®¤à¯à®¤à®®à¯", "â‚¹21,300 - â‚¹31,000", "Deposit à®‰à®Ÿà¯à®ªà®Ÿ approximate setup range"]], [55*mm, 42*mm, 79*mm], style))

    story.append(PageBreak())
    story.append(P("8. Advance à®•à¯Šà®Ÿà¯à®•à¯à®•à¯à®®à¯ à®®à¯à®©à¯ verify à®šà¯†à®¯à¯à®¯ à®µà¯‡à®£à¯à®Ÿà®¿à®¯à®µà¯ˆ", style["h1"]))
    for item in VERIFY_ITEMS:
        story.append(P("â€¢ " + item, style["body"]))

    story.append(P("9. Independent room-à®©à¯ à®¨à®©à¯à®®à¯ˆà®•à®³à¯", style["h1"]))
    story.append(P("Independent room hostel-à® à®µà®¿à®Ÿ à®Žà®©à®•à¯à®•à¯ productivity-à®•à¯à®•à¯ à®à®±à¯à®±à®¤à¯. Better focus, privacy, health control, software projects, AI learning, MBA preparation à®†à®•à®¿à®¯à®µà®±à¯à®±à®¿à®²à¯ à®¨à¯‡à®°à®¤à¯à®¤à¯ˆ à®šà¯‡à®®à®¿à®¤à¯à®¤à¯ à®¨à®¿à®²à¯ˆà®¯à®¾à®© à®®à¯à®©à¯à®©à¯‡à®±à¯à®±à®®à¯ à®ªà¯†à®± à®‰à®¤à®µà¯à®®à¯. Hostel-à®²à¯ à®•à®¿à®Ÿà¯ˆà®•à¯à®•à¯à®®à¯ social environment à®¨à®²à¯à®²à®¤à¯‡; à®†à®©à®¾à®²à¯ à®‡à®¨à¯à®¤ à®’à®°à¯ à®µà®°à¯à®Ÿà®¤à¯à®¤à®¿à®²à¯ à®Žà®©à®•à¯à®•à¯ à®¤à¯‡à®µà¯ˆà®¯à®¾à®©à®¤à¯ à®…à®®à¯ˆà®¤à®¿à®¯à®¾à®© execution environment.", style["body"]))

    story.append(P("10. Possible Risks à®®à®±à¯à®±à¯à®®à¯ à®®à¯‡à®²à®¾à®£à¯à®®à¯ˆ", style["h1"]))
    story.append(make_table([["Risk", "Manage à®šà¯†à®¯à¯à®¯à¯à®®à¯ à®®à¯à®±à¯ˆ"]] + RISKS, [52*mm, 124*mm], style))

    story.append(P("11. à®‡à®±à¯à®¤à®¿ à®ªà®°à®¿à®¨à¯à®¤à¯à®°à¯ˆ", style["h1"]))
    conclusion = (
        "à®…à®ªà¯à®ªà®¾, à®‡à®¨à¯à®¤ à®®à¯à®Ÿà®¿à®µà¯ˆ à®¨à®¾à®©à¯ à®‰à®£à®°à¯à®šà¯à®šà®¿à®µà®šà®ªà¯à®ªà®Ÿà¯à®Ÿà¯ à®Žà®Ÿà¯à®•à¯à®•à®µà®¿à®²à¯à®²à¯ˆ. à®•à®²à¯à®²à¯‚à®°à®¿à®•à¯à®•à¯ à®¤à¯‚à®°à®®à¯, à®µà®¾à®Ÿà®•à¯ˆ, deposit, "
        "à®ªà®¾à®¤à¯à®•à®¾à®ªà¯à®ªà¯, internet, à®¤à®£à¯à®£à¯€à®°à¯, future MBA plan à®†à®•à®¿à®¯à®µà®±à¯à®±à¯ˆ à®µà¯ˆà®¤à¯à®¤à¯ à®•à®µà®©à®®à®¾à®- à®†à®¯à¯à®µà¯ à®šà¯†à®¯à¯à®¤à¯‡à®©à¯. "
        "à®‡à®¨à¯à®¤ à®µà¯€à®Ÿà¯ à®šà¯à®®à®¾à®°à¯ à®’à®°à¯ à®µà®°à¯à®Ÿà®¤à¯à®¤à®¿à®±à¯à®•à¯ à®®à®Ÿà¯à®Ÿà¯à®®à¯. Graduation à®®à¯à®Ÿà®¿à®¨à¯à®¤ à®ªà®¿à®±à®•à¯ MBA abroad à®¨à¯‹à®•à¯à®•à®¿ à®¨à®•à®°à¯à®®à¯ à®¤à®¿à®Ÿà¯à®Ÿà®®à¯ à®‰à®³à¯à®³à®¤à¯. "
        "à®…à®¤à®©à®¾à®²à¯ à®‡à®¤à¯ à®¨à®¿à®°à®¨à¯à®¤à®° à®šà¯†à®²à®µà¯ à®…à®²à¯à®²; better productivity, health, projects, preparation à®†à®•à®¿à®¯à®µà®±à¯à®±à¯à®•à¯à®•à®¾à®© temporary investment. "
        "à®¨à¯€à®™à¯à®•à®³à¯ à®ªà®¾à®°à¯à®¤à¯à®¤à¯ à®‰à®™à¯à®•à®³à¯ à®•à®°à¯à®¤à¯à®¤à¯ˆ à®šà¯Šà®²à¯à®²à¯à®™à¯à®•à®³à¯. à®¨à¯€à®™à¯à®•à®³à¯ à®šà®°à®¿ à®Žà®©à¯à®±à¯ à®¨à®¿à®©à¯ˆà®¤à¯à®¤ à®ªà®¿à®±à®•à¯‡ advance à®•à¯Šà®Ÿà¯à®ªà¯à®ªà¯‹à®®à¯."
    )
    add_card(story, "à®®à®°à®¿à®¯à®¾à®¤à¯ˆà®¯à®¾à®© à®µà¯‡à®£à¯à®Ÿà¯à®•à¯‹à®³à¯", [conclusion], style)

    def footer(canvas, doc_):
        canvas.saveState()
        canvas.setFont("NotoTamil", 7)
        canvas.setFillColor(MUTED)
        canvas.drawString(16 * mm, 8 * mm, "à®°à®¤à¯à®¤à®¿à®©à®®à¯ à®•à®²à¯à®²à¯‚à®°à®¿ à®…à®°à¯à®•à®¿à®²à¯à®³à¯à®³ à®µà®¾à®Ÿà®•à¯ˆ à®µà¯€à®Ÿà¯ - à®†à®¯à¯à®µà¯ à®…à®±à®¿à®•à¯à®•à¯ˆ")
        canvas.drawRightString(194 * mm, 8 * mm, f"Page {doc_.page}")
        canvas.restoreState()

    doc.build(story, onFirstPage=footer, onLaterPages=footer)
    return path


def add_docx_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color="111827", size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(str(text))
    r.font.name = "Noto Sans Tamil"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans Tamil")
    r._element.rPr.rFonts.set(qn("w:cs"), "Noto Sans Tamil")
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = RGBColor.from_string(color)


def add_docx_table(doc, rows, widths=None):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            set_cell_text(cell, value, bold=(r_idx == 0), color=("FFFFFF" if r_idx == 0 else "111827"), size=8)
            if r_idx == 0:
                add_docx_shading(cell, "123A63")
            elif r_idx % 2 == 0:
                add_docx_shading(cell, "F5F9FE")
    doc.add_paragraph()
    return table


def docx_para(doc, text, style=None, bold=False, size=10.5, color="111827", align=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.font.name = "Noto Sans Tamil"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans Tamil")
    r._element.rPr.rFonts.set(qn("w:cs"), "Noto Sans Tamil")
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = RGBColor.from_string(color)
    p.paragraph_format.space_after = Pt(5)
    return p


def build_docx(houses):
    info = tamil_content(houses)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)

    styles = doc.styles
    styles["Normal"].font.name = "Noto Sans Tamil"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans Tamil")
    styles["Normal"]._element.rPr.rFonts.set(qn("w:cs"), "Noto Sans Tamil")
    styles["Normal"].font.size = Pt(10.5)

    header = section.header.paragraphs[0]
    header.text = "à®°à®¤à¯à®¤à®¿à®©à®®à¯ à®•à®²à¯à®²à¯‚à®°à®¿ à®…à®°à¯à®•à®¿à®²à¯à®³à¯à®³ à®µà®¾à®Ÿà®•à¯ˆ à®µà¯€à®Ÿà¯"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer = section.footer.paragraphs[0]
    footer.text = "Prepared by Dragon"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    docx_para(doc, "à®°à®¤à¯à®¤à®¿à®©à®®à¯ à®•à®²à¯à®²à¯‚à®°à®¿ à®…à®°à¯à®•à®¿à®²à¯à®³à¯à®³ à®µà®¾à®Ÿà®•à¯ˆ à®µà¯€à®Ÿà¯", bold=True, size=22, color="123A63", align=WD_ALIGN_PARAGRAPH.CENTER)
    docx_para(doc, "à®®à¯à®´à¯à®®à¯ˆà®¯à®¾à®© à®†à®¯à¯à®µà¯ à®…à®±à®¿à®•à¯à®•à¯ˆ", bold=True, size=18, color="123A63", align=WD_ALIGN_PARAGRAPH.CENTER)
    docx_para(doc, "Prepared for: My Father\nPrepared by: Dragon\nDate: " + TODAY_TAMIL, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    sections = [
        ("1. Hostel-à®²à®¿à®°à¯à®¨à¯à®¤à¯ à®µà¯†à®³à®¿à®¯à¯‡à®± à®µà®¿à®°à¯à®®à¯à®ªà¯à®®à¯ à®•à®¾à®°à®£à®®à¯", [info["intro"],
         "Better concentration, quiet environment, better sleep, better food control, AI development, software projects, MBA preparation, privacy, independence à®†à®•à®¿à®¯à®µà¯ˆ à®‡à®¨à¯à®¤ à®®à®¾à®±à¯à®±à®¤à¯à®¤à®¿à®©à¯ à®®à¯à®•à¯à®•à®¿à®¯ à®•à®¾à®°à®£à®™à¯à®•à®³à¯. à®‡à®¤à¯ à®šà¯à®®à®¾à®°à¯ à®’à®°à¯ à®µà®°à¯à®Ÿà®¤à¯à®¤à®¿à®±à¯à®•à¯ à®®à®Ÿà¯à®Ÿà¯à®®à¯."]),
        ("2. à®Žà®©à¯ à®¤à¯‡à®µà¯ˆà®•à®³à¯", ["â‚¹5,000-â‚¹10,000 rent, Independent room / 1 RK / 1 BHK, No PG, No Hostel, safe area, peaceful neighbourhood, hospital/grocery à®…à®°à¯à®•à®¿à®²à¯, bike travel 15-20 minutes-à®•à¯à®•à¯ à®•à¯€à®´à¯, à®¨à®²à¯à®² internet/water, easy move-out."]),
    ]
    for title, paras in sections:
        docx_para(doc, title, bold=True, size=15, color="123A63")
        for para in paras:
            docx_para(doc, para)

    docx_para(doc, "3. à®ªà®•à¯à®¤à®¿à®•à®³à¯ à®’à®ªà¯à®ªà¯€à®Ÿà¯", bold=True, size=15, color="123A63")
    add_docx_table(doc, area_table_rows())
    docx_para(doc, "4. Top 10 à®µà¯€à®Ÿà¯à®•à®³à¯", bold=True, size=15, color="123A63")
    add_docx_table(doc, house_table_rows(houses))
    doc.add_page_break()

    best = info["best"]
    docx_para(doc, "5. à®ªà®°à®¿à®¨à¯à®¤à¯à®°à¯ˆà®•à¯à®•à®ªà¯à®ªà®Ÿà¯à®®à¯ à®µà¯€à®Ÿà¯", bold=True, size=15, color="123A63")
    docx_para(doc, f"Recommended House: {best['title']} - {best['address']}", bold=True, size=12, color="2E74B5")
    docx_para(doc, f"Rent: {rupee(best['rent'])}; Security Deposit: {rupee(best['deposit'])}; Bike distance: {best['distance']} km; Travel time: {best['time']}; Furnishing: {best['furnished']}.")
    docx_para(doc, "à®‡à®¨à¯à®¤ à®µà¯€à®Ÿà¯ à®•à¯à®±à¯ˆà®¨à¯à®¤ à®šà¯†à®²à®µà¯, à®•à¯à®±à¯ˆà®¨à¯à®¤ commute, Semi-Furnished convenience, temporary one-year stay à®†à®•à®¿à®¯à®µà®±à¯à®±à¯à®•à¯à®•à¯ à®¨à®²à¯à®² balance à®¤à®°à¯à®•à®¿à®±à®¤à¯.")

    docx_para(doc, "6. à®Žà®¤à®¿à®°à¯à®ªà®¾à®°à¯à®•à¯à®•à®ªà¯à®ªà®Ÿà¯à®®à¯ à®®à®¾à®¤à®šà¯ à®šà¯†à®²à®µà¯", bold=True, size=15, color="123A63")
    budget_rows = [["à®šà¯†à®²à®µà¯", "à®¤à®±à¯à®ªà¯‹à®¤à¯ˆà®¯ à®®à®¤à®¿à®ªà¯à®ªà¯", "à®•à¯à®±à¯ˆà®¨à¯à®¤à®ªà®Ÿà¯à®šà®®à¯", "à®…à®¤à®¿à®•à®ªà®Ÿà¯à®šà®®à¯"]]
    for name, current, minimum, maximum in MONTHLY_BUDGET:
        budget_rows.append([name, rupee(current), rupee(minimum), rupee(maximum)])
    budget_rows.append(["à®®à¯Šà®¤à¯à®¤à®®à¯", rupee(info["expected_monthly"]), rupee(info["min_monthly"]), rupee(info["max_monthly"])])
    add_docx_table(doc, budget_rows)
    docx_para(doc, f"à®’à®°à¯ à®µà®°à¯à®Ÿ à®šà¯†à®²à®µà¯ à®•à®£à®¿à®ªà¯à®ªà¯: à®šà¯à®®à®¾à®°à¯ {rupee(info['yearly'])}.")

    docx_para(doc, "7. One-Time Expenses", bold=True, size=15, color="123A63")
    add_docx_table(doc, [["Item", "Approx. Amount", "à®•à¯à®±à®¿à®ªà¯à®ªà¯"]] + ONE_TIME + [["à®®à¯Šà®¤à¯à®¤à®®à¯", "â‚¹21,300 - â‚¹31,000", "Deposit à®‰à®Ÿà¯à®ªà®Ÿ approximate setup range"]])

    docx_para(doc, "8. Advance à®•à¯Šà®Ÿà¯à®•à¯à®•à¯à®®à¯ à®®à¯à®©à¯ verify à®šà¯†à®¯à¯à®¯ à®µà¯‡à®£à¯à®Ÿà®¿à®¯à®µà¯ˆ", bold=True, size=15, color="123A63")
    for item in VERIFY_ITEMS:
        docx_para(doc, "â€¢ " + item)

    docx_para(doc, "9. Independent room-à®©à¯ à®¨à®©à¯à®®à¯ˆà®•à®³à¯", bold=True, size=15, color="123A63")
    docx_para(doc, "Independent room hostel-à® à®µà®¿à®Ÿ better focus, privacy, health, projects, MBA preparation, time saving à®†à®•à®¿à®¯à®µà®±à¯à®±à¯à®•à¯à®•à¯ à®‰à®¤à®µà¯à®®à¯.")
    docx_para(doc, "10. Possible Risks à®®à®±à¯à®±à¯à®®à¯ à®®à¯‡à®²à®¾à®£à¯à®®à¯ˆ", bold=True, size=15, color="123A63")
    add_docx_table(doc, [["Risk", "Manage à®šà¯†à®¯à¯à®¯à¯à®®à¯ à®®à¯à®±à¯ˆ"]] + RISKS)
    docx_para(doc, "11. à®‡à®±à¯à®¤à®¿ à®ªà®°à®¿à®¨à¯à®¤à¯à®°à¯ˆ", bold=True, size=15, color="123A63")
    docx_para(doc, "à®…à®ªà¯à®ªà®¾, à®‡à®¨à¯à®¤ à®®à¯à®Ÿà®¿à®µà¯ˆ à®¨à®¾à®©à¯ à®‰à®£à®°à¯à®šà¯à®šà®¿à®µà®šà®ªà¯à®ªà®Ÿà¯à®Ÿà¯ à®Žà®Ÿà¯à®•à¯à®•à®µà®¿à®²à¯à®²à¯ˆ. à®•à®²à¯à®²à¯‚à®°à®¿à®•à¯à®•à¯ à®¤à¯‚à®°à®®à¯, à®µà®¾à®Ÿà®•à¯ˆ, deposit, à®ªà®¾à®¤à¯à®•à®¾à®ªà¯à®ªà¯, internet, à®¤à®£à¯à®£à¯€à®°à¯, future MBA plan à®†à®•à®¿à®¯à®µà®±à¯à®±à¯ˆ à®µà¯ˆà®¤à¯à®¤à¯ à®•à®µà®©à®®à®¾à®- à®†à®¯à¯à®µà¯ à®šà¯†à®¯à¯à®¤à¯‡à®©à¯. à®‡à®¤à¯ à®šà¯à®®à®¾à®°à¯ à®’à®°à¯ à®µà®°à¯à®Ÿà®¤à¯à®¤à®¿à®±à¯à®•à¯ à®®à®Ÿà¯à®Ÿà¯à®®à¯. à®¨à¯€à®™à¯à®•à®³à¯ à®ªà®¾à®°à¯à®¤à¯à®¤à¯ à®‰à®™à¯à®•à®³à¯ à®•à®°à¯à®¤à¯à®¤à¯ˆ à®šà¯Šà®²à¯à®²à¯à®™à¯à®•à®³à¯; à®¨à¯€à®™à¯à®•à®³à¯ à®šà®°à®¿ à®Žà®©à¯à®±à¯ à®¨à®¿à®©à¯ˆà®¤à¯à®¤ à®ªà®¿à®±à®•à¯‡ advance à®•à¯Šà®Ÿà¯à®ªà¯à®ªà¯‹à®®à¯.")

    path = os.path.join(OUTDIR, "rathinam_house_report_tamil.docx")
    doc.save(path)
    return path


def build_markdown(houses):
    info = tamil_content(houses)
    lines = [
        "# à®°à®¤à¯à®¤à®¿à®©à®®à¯ à®•à®²à¯à®²à¯‚à®°à®¿ à®…à®°à¯à®•à®¿à®²à¯à®³à¯à®³ à®µà®¾à®Ÿà®•à¯ˆ à®µà¯€à®Ÿà¯",
        "## à®®à¯à®´à¯à®®à¯ˆà®¯à®¾à®© à®†à®¯à¯à®µà¯ à®…à®±à®¿à®•à¯à®•à¯ˆ",
        "",
        "Prepared for: My Father  ",
        "Prepared by: Dragon  ",
        f"Date: {TODAY_TAMIL}",
        "",
        "## à®®à¯à®•à¯à®•à®¿à®¯ à®®à¯à®Ÿà®¿à®µà¯",
        f"Recommended house: {info['best']['title']} - {info['best']['address']}. Rent {rupee(info['best']['rent'])}, Security Deposit {rupee(info['best']['deposit'])}, distance {info['best']['distance']} km, travel time {info['best']['time']}.",
        "",
        "## 1. Hostel-à®²à®¿à®°à¯à®¨à¯à®¤à¯ à®µà¯†à®³à®¿à®¯à¯‡à®± à®µà®¿à®°à¯à®®à¯à®ªà¯à®®à¯ à®•à®¾à®°à®£à®®à¯",
        info["intro"],
        "",
        "- Better concentration",
        "- Quiet environment",
        "- Better sleep",
        "- Better food control",
        "- AI development à®®à®±à¯à®±à¯à®®à¯ software projects",
        "- MBA preparation",
        "- Privacy à®®à®±à¯à®±à¯à®®à¯ independence",
        "- à®‡à®¤à¯ à®šà¯à®®à®¾à®°à¯ à®’à®°à¯ à®µà®°à¯à®Ÿà®¤à¯à®¤à®¿à®±à¯à®•à¯ à®®à®Ÿà¯à®Ÿà¯à®®à¯",
        "",
        "## 2. à®Žà®©à¯ à®¤à¯‡à®µà¯ˆà®•à®³à¯",
    ]
    for item in [
        "â‚¹5,000-â‚¹10,000 rent",
        "Independent room / 1 RK / 1 BHK",
        "No PG, No Hostel",
        "Safe and peaceful area",
        "Hospital/grocery à®…à®°à¯à®•à®¿à®²à¯",
        "Bike travel below 15-20 minutes",
        "Good internet and water",
        "Easy move-out after one year",
    ]:
        lines.append(f"- {item}")
    lines += ["", "## 3. à®ªà®•à¯à®¤à®¿à®•à®³à¯ à®’à®ªà¯à®ªà¯€à®Ÿà¯", ""]
    lines.append("| " + " | ".join(map(str, area_table_rows()[0])) + " |")
    lines.append("|" + "|".join(["---"] * len(area_table_rows()[0])) + "|")
    for row in area_table_rows()[1:]:
        lines.append("| " + " | ".join(map(str, row)) + " |")
    lines += ["", "## 4. Top 10 à®µà¯€à®Ÿà¯à®•à®³à¯", ""]
    house_rows = house_table_rows(houses)
    lines.append("| " + " | ".join(house_rows[0]) + " |")
    lines.append("|" + "|".join(["---"] * len(house_rows[0])) + "|")
    for row in house_rows[1:]:
        lines.append("| " + " | ".join(map(str, row)) + " |")
    lines += [
        "",
        "## 5. à®ªà®°à®¿à®¨à¯à®¤à¯à®°à¯ˆà®•à¯à®•à®ªà¯à®ªà®Ÿà¯à®®à¯ à®µà¯€à®Ÿà¯",
        f"{info['best']['title']} - {info['best']['address']}. à®‡à®¤à¯ à®•à¯à®±à¯ˆà®¨à¯à®¤ rent, à®•à¯à®±à¯ˆà®¨à¯à®¤ commute, Semi-Furnished à®µà®šà®¤à®¿, temporary stay suitability à®†à®•à®¿à®¯à®µà®±à¯à®±à®¾à®²à¯ à®šà®¿à®±à®¨à¯à®¤ balance à®¤à®°à¯à®•à®¿à®±à®¤à¯.",
        "",
        "## 6. Expected Expenses",
        "| à®šà¯†à®²à®µà¯ | à®¤à®±à¯à®ªà¯‹à®¤à¯ˆà®¯ à®®à®¤à®¿à®ªà¯à®ªà¯ | à®•à¯à®±à¯ˆà®¨à¯à®¤à®ªà®Ÿà¯à®šà®®à¯ | à®…à®¤à®¿à®•à®ªà®Ÿà¯à®šà®®à¯ |",
        "|---|---:|---:|---:|",
    ]
    for name, current, minimum, maximum in MONTHLY_BUDGET:
        lines.append(f"| {name} | {rupee(current)} | {rupee(minimum)} | {rupee(maximum)} |")
    lines.append(f"| à®®à¯Šà®¤à¯à®¤à®®à¯ | {rupee(info['expected_monthly'])} | {rupee(info['min_monthly'])} | {rupee(info['max_monthly'])} |")
    lines += ["", f"Yearly cost: à®šà¯à®®à®¾à®°à¯ {rupee(info['yearly'])}.", "", "## 7. One-Time Expenses"]
    for row in ONE_TIME:
        lines.append(f"- {row[0]}: {row[1]} - {row[2]}")
    lines += ["", "## 8. Advance à®•à¯Šà®Ÿà¯à®•à¯à®•à¯à®®à¯ à®®à¯à®©à¯ verify à®šà¯†à®¯à¯à®¯ à®µà¯‡à®£à¯à®Ÿà®¿à®¯à®µà¯ˆ"]
    for item in VERIFY_ITEMS:
        lines.append(f"- {item}")
    lines += ["", "## 9. Advantages", "Independent room à®Žà®©à®•à¯à®•à¯ focus, privacy, health, projects, MBA preparation, time saving à®†à®•à®¿à®¯à®µà®±à¯à®±à¯à®•à¯à®•à¯ hostel-à® à®µà®¿à®Ÿ à®¨à®²à¯à®²à®¤à¯.", "", "## 10. Possible Risks"]
    for risk, manage in RISKS:
        lines.append(f"- **{risk}:** {manage}")
    lines += ["", "## 11. à®‡à®±à¯à®¤à®¿ à®ªà®°à®¿à®¨à¯à®¤à¯à®°à¯ˆ", "à®…à®ªà¯à®ªà®¾, à®‡à®¨à¯à®¤ à®®à¯à®Ÿà®¿à®µà¯ˆ à®¨à®¾à®©à¯ à®‰à®£à®°à¯à®šà¯à®šà®¿à®µà®šà®ªà¯à®ªà®Ÿà¯à®Ÿà¯ à®Žà®Ÿà¯à®•à¯à®•à®µà®¿à®²à¯à®²à¯ˆ. à®‡à®¤à¯ à®šà¯à®®à®¾à®°à¯ à®’à®°à¯ à®µà®°à¯à®Ÿà®¤à¯à®¤à®¿à®±à¯à®•à¯ à®®à®Ÿà¯à®Ÿà¯à®®à¯; graduation-à®•à¯à®•à¯à®ªà¯ à®ªà®¿à®±à®•à¯ MBA abroad plan à®‰à®³à¯à®³à®¤à¯. à®‰à®™à¯à®•à®³à¯ à®•à®°à¯à®¤à¯à®¤à¯ˆ à®®à®°à®¿à®¯à®¾à®¤à¯ˆà®¯à¯à®Ÿà®©à¯ à®•à¯‡à®Ÿà¯à®•à®¿à®±à¯‡à®©à¯."]
    path = os.path.join(OUTDIR, "rathinam_house_report_tamil.md")
    with open(path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    return path


def build_whatsapp(houses):
    info = tamil_content(houses)
    text = f"""à®°à®¤à¯à®¤à®¿à®©à®®à¯ à®•à®²à¯à®²à¯‚à®°à®¿ à®…à®°à¯à®•à®¿à®²à¯ à®µà®¾à®Ÿà®•à¯ˆ à®µà¯€à®Ÿà¯ - à®šà¯à®°à¯à®•à¯à®•à®®à¯

à®…à®ªà¯à®ªà®¾, à®‡à®¤à¯ à®šà¯à®®à®¾à®°à¯ 10-12 à®®à®¾à®¤à®™à¯à®•à®³à¯à®•à¯à®•à¯ à®®à®Ÿà¯à®Ÿà¯à®®à¯ temporary stay.

Recommended:
{info['best']['title']}
Location: {info['best']['address']}
Rent: {rupee(info['best']['rent'])}
Security Deposit: {rupee(info['best']['deposit'])}
Distance: {info['best']['distance']} km
Bike time: {info['best']['time']}
Furnishing: {info['best']['furnished']}

à®à®©à¯ à®‡à®¤à¯ à®¨à®²à¯à®²à®¤à¯:
- à®•à®²à¯à®²à¯‚à®°à®¿à®•à¯à®•à¯ à®®à®¿à®- à®…à®°à¯à®•à®¿à®²à¯
- à®•à¯à®±à¯ˆà®¨à¯à®¤ rent + à®•à¯à®±à¯ˆà®¨à¯à®¤ deposit
- Semi-Furnished; setup à®šà¯†à®²à®µà¯ à®•à¯à®±à¯ˆà®¯à¯à®®à¯
- AI projects, software work, MBA preparation-à®•à¯à®•à¯ à®…à®®à¯ˆà®¤à®¿à®¯à®¾à®© à®šà¯‚à®´à®²à¯
- Hostel-à® à®µà®¿à®Ÿ privacy, focus, sleep, food control à®®à¯‡à®®à¯à®ªà®Ÿà¯à®®à¯

Monthly expected cost: {rupee(info['expected_monthly'])}
One-year estimate: {rupee(info['yearly'])}

Advance à®•à¯Šà®Ÿà¯à®•à¯à®•à¯à®®à¯ à®®à¯à®©à¯ verify:
Water, Fiber Internet, parking, neighbours, night safety, rent agreement, deposit refund, power cuts, EB meter, mobile signal.

à®‡à®¤à¯ emotional decision à®‡à®²à¯à®²à¯ˆ. Carefully researched temporary productivity investment. à®¨à¯€à®™à¯à®•à®³à¯ à®ªà®¾à®°à¯à®¤à¯à®¤à¯ à®‰à®™à¯à®•à®³à¯ à®•à®°à¯à®¤à¯à®¤à¯ à®šà¯Šà®²à¯à®²à¯à®™à¯à®•à®³à¯."""
    path = os.path.join(OUTDIR, "whatsapp_one_page_summary_tamil.txt")
    with open(path, "w", encoding="utf-8") as f:
        f.write(text)
    return path


if __name__ == "__main__":
    houses = load_data()
    outputs = {
        "pdf": build_pdf(houses),
        "docx": build_docx(houses),
        "markdown": build_markdown(houses),
        "whatsapp": build_whatsapp(houses),
    }
    for key, value in outputs.items():
        print(f"{key}: {value}")

