import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
OUT = DOCS / "generated"
OUT.mkdir(exist_ok=True)
PROFILE_EXAMPLE = ROOT / "profile_data.example.json"
PROFILE_LOCAL = ROOT / "profile_data.json"


DEFAULT_PROFILE = {
    "name": "Deepan Karthick",
    "email": "ADD_EMAIL",
    "phone": "ADD_PHONE",
    "city_country": "ADD_CITY_COUNTRY",
    "github_url": "ADD_GITHUB_URL",
    "linkedin_url": "ADD_LINKEDIN_URL",
    "portfolio_url": "ADD_PORTFOLIO_URL",
    "degree": "ADD_DEGREE",
    "college": "ADD_COLLEGE",
    "graduation_year": "ADD_GRADUATION_YEAR",
    "experience_dates": "ADD_DATES",
}


def load_profile():
    source = PROFILE_LOCAL if PROFILE_LOCAL.exists() else PROFILE_EXAMPLE
    if source.exists():
        with source.open("r", encoding="utf-8") as fh:
            data = json.load(fh)
        return {**DEFAULT_PROFILE, **data}
    return DEFAULT_PROFILE.copy()


def contact_line(profile):
    return (
        f"{profile['email']} | {profile['phone']} | {profile['city_country']} | "
        f"GitHub: {profile['github_url']} | LinkedIn: {profile['linkedin_url']} | "
        f"Portfolio: {profile['portfolio_url']}"
    )


def render_value(text, profile):
    replacements = {
        "ADD_EMAIL": profile["email"],
        "ADD_PHONE": profile["phone"],
        "ADD_CITY_COUNTRY": profile["city_country"],
        "ADD_GITHUB_URL": profile["github_url"],
        "ADD_LINKEDIN_URL": profile["linkedin_url"],
        "ADD_PORTFOLIO_URL": profile["portfolio_url"],
        "ADD_DEGREE": profile["degree"],
        "ADD_COLLEGE": profile["college"],
        "ADD_GRADUATION_YEAR": profile["graduation_year"],
        "ADD_DATES": profile["experience_dates"],
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text


ATS = {
    "filename": "karthik_ats_resume.docx",
    "pdf_filename": "karthik_ats_resume.pdf",
    "subtitle": "B.Sc. CS (AI) Student | AI Automation Developer | React | Python",
    "sections": [
        (
            "Summary",
            [
                "Computer Science with Artificial Intelligence student building Python automation tools, React dashboards, ERP-style workflows, and PDF/DOCX report generators. Strongest proof: rental research automation with filtering/scoring logic, offline sample data, generated reports, and unit tests."
            ],
        ),
        (
            "Skills",
            [
                "Languages: Python, JavaScript, HTML, CSS, TypeScript basics",
                "Frontend and data: React, Vite, Supabase, ReportLab, python-docx, Django basics",
                "Tools and areas: Git, GitHub, GitHub Actions, VS Code, CLI tools, AI-assisted development with manual verification, REST API basics, PostgreSQL basics, ERP workflows, dashboards, testing, documentation, Docker basics, GSAP basics, Framer Motion basics",
            ],
        ),
        (
                "Projects",
            [
                "Rental Research Report Generator | Python, ReportLab, python-docx, JSON",
                "Built a Python automation workflow that normalizes rental listing data, filters unsuitable options, ranks matches, and generates PDF/DOCX decision reports.",
                "Implemented scoring/filters for rent range, listing type, commute distance, furnishing, deposit practicality, and unsuitable categories.",
                "Generated English/Tamil outputs with ranked properties, budgets, verification checklist, source links, offline sample data, and unit tests.",
                "Printing Press ERP / Business Management System | React, Supabase, JavaScript, Vite",
                "Built and deployed React/Supabase business software screens for printing press workflows, dashboards, and operational tracking.",
                "Modeled local-business workflows including inventory, production tracking, order operations, and role-based access concepts.",
                "Portfolio Website | React, Vite, Framer Motion, GSAP",
                "Created a public developer portfolio and one-tap links page with project proof, resume links, GitHub, LinkedIn, live demos, and contact actions.",
            ],
        ),
        (
            "Education",
            [
                "ADD_DEGREE, ADD_COLLEGE - ADD_GRADUATION_YEAR",
                "Relevant coursework: Artificial Intelligence, Programming, Data Structures, Web Development, Database Systems",
            ],
        ),
        (
            "Additional",
            [
                "Availability: Immediate; remote internships, freelance work, part-time developer roles, contract projects around college schedule",
                "Languages: English, Tamil",
            ],
        ),
    ],
}


STARTUP = {
    "filename": "karthik_startup_resume.docx",
    "pdf_filename": "karthik_startup_resume.pdf",
    "subtitle": "Python automation + React dashboards + practical business software",
    "sections": [
        (
            "Profile",
            [
                "I build practical tools that turn messy work into usable software: ranked lists, PDF/DOCX reports, React dashboards, ERP-style workflows, automations, and AI-assisted development workflows with manual verification. I am looking for remote internships, freelance projects, contract work, and part-time developer roles where speed, ownership, and clear communication matter."
            ],
        ),
        (
            "What I Can Help With",
            [
                "Python scripts that automate repetitive workflows",
                "PDF/DOCX report generation",
                "React dashboards and internal-tool screens",
                "Data cleanup and structured summaries",
                "Internal tools for small teams",
                "AI-assisted prototyping, prompt workflows, documentation, and testing",
            ],
        ),
        (
                "Best Proof",
            [
                "Rental Research Report Generator | Python, ReportLab, python-docx, JSON",
                "Built a Python tool that normalizes rental listing data, filters noisy results, ranks suitable houses, and generates decision reports in English and Tamil.",
                "Replaced scattered rental searching with a structured shortlist.",
                "Turned public listing data into a ranked report with budget and commute logic.",
                "Produced family-ready outputs: PDF, DOCX, Markdown, and WhatsApp summary.",
                "Included real-world safeguards: water, internet, safety, deposit, and agreement verification checklist.",
                "Added offline sample data and unit tests so the project can be shown reliably without live portals.",
                "Printing Press ERP / Business Management System | React, Supabase, business workflows",
                "Built and deployed practical React/Supabase business software for workflow tracking and operational dashboards.",
            ],
        ),
        (
            "Skills",
            [
                "Python, JavaScript, React, Vite, Supabase, Django basics, HTML, CSS, TypeScript basics, REST APIs, PostgreSQL, Git, GitHub, GitHub Actions, Docker basics, Codex, ChatGPT, AI-assisted development, prompt engineering, ERP development, dashboard development, business software, automation, PDF/DOCX generation, responsive web design, UI/UX design, basic testing, CI/CD basics, CLI tools, technical documentation, GSAP basics, Framer Motion basics"
            ],
        ),
        (
            "Work Style",
            [
                "I communicate progress clearly.",
                "I document setup and usage.",
                "I prefer small working versions over vague big promises.",
                "I verify AI-assisted work before treating it as finished.",
                "I can learn fast when the problem is real.",
            ],
        ),
        (
            "Education",
            ["ADD_DEGREE, ADD_COLLEGE - ADD_GRADUATION_YEAR"],
        ),
    ],
}


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color in [
        ("Heading 1", 13, "0F766E"),
        ("Heading 2", 11, "111827"),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(3)


def add_para(doc, text, bold=False, size=10, color="111827", align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(text)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style=None)
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("- " + text)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string("111827")


def build_resume(data, profile):
    doc = Document()
    style_doc(doc)

    add_para(doc, profile["name"], bold=True, size=20, color="111827", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, data["subtitle"], size=10.5, color="0F766E", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(
        doc,
        contact_line(profile),
        size=8.8,
        color="4B5563",
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    for title, items in data["sections"]:
        doc.add_paragraph(title, style="Heading 1")
        if title in {"Summary", "Profile"}:
            for item in items:
                add_para(doc, render_value(item, profile))
        elif title == "Skills" and len(items) == 1:
            add_para(doc, render_value(items[0], profile))
        else:
            for item in items:
                item = render_value(item, profile)
                if "|" in item:
                    add_para(doc, item, bold=True, color="111827")
                else:
                    add_bullet(doc, item)

    path = OUT / data["filename"]
    doc.save(path)
    return path


def pdf_styles():
    base = getSampleStyleSheet()
    return {
        "name": ParagraphStyle(
            "Name",
            parent=base["Title"],
            fontName="Helvetica-Bold",
            fontSize=20,
            leading=24,
            alignment=1,
            textColor=colors.HexColor("#111827"),
            spaceAfter=3,
        ),
        "subtitle": ParagraphStyle(
            "Subtitle",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=10,
            leading=13,
            alignment=1,
            textColor=colors.HexColor("#0F766E"),
            spaceAfter=3,
        ),
        "contact": ParagraphStyle(
            "Contact",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=8,
            leading=10,
            alignment=1,
            textColor=colors.HexColor("#4B5563"),
            spaceAfter=8,
        ),
        "section": ParagraphStyle(
            "Section",
            parent=base["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=12,
            leading=15,
            textColor=colors.HexColor("#0F766E"),
            spaceBefore=8,
            spaceAfter=3,
        ),
        "body": ParagraphStyle(
            "Body",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=9.2,
            leading=12,
            textColor=colors.HexColor("#111827"),
            spaceAfter=3,
        ),
        "strong": ParagraphStyle(
            "Strong",
            parent=base["BodyText"],
            fontName="Helvetica-Bold",
            fontSize=9.2,
            leading=12,
            textColor=colors.HexColor("#111827"),
            spaceAfter=3,
        ),
        "bullet": ParagraphStyle(
            "Bullet",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=9.0,
            leading=11.5,
            leftIndent=12,
            firstLineIndent=-8,
            textColor=colors.HexColor("#111827"),
            spaceAfter=2,
        ),
    }


def xml_escape(text):
    return (
        str(text)
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def build_resume_pdf(data, profile):
    styles = pdf_styles()
    path = OUT / data["pdf_filename"]
    doc = SimpleDocTemplate(
        str(path),
        pagesize=LETTER,
        leftMargin=0.65 * inch,
        rightMargin=0.65 * inch,
        topMargin=0.58 * inch,
        bottomMargin=0.58 * inch,
    )
    story = [
        Paragraph(xml_escape(profile["name"]), styles["name"]),
        Paragraph(xml_escape(data["subtitle"]), styles["subtitle"]),
        Paragraph(xml_escape(contact_line(profile)), styles["contact"]),
    ]

    for title, items in data["sections"]:
        story.append(Paragraph(xml_escape(title), styles["section"]))
        for item in items:
            item = render_value(item, profile)
            if title in {"Summary", "Profile"}:
                story.append(Paragraph(xml_escape(item), styles["body"]))
            elif title == "Skills" and len(items) == 1:
                story.append(Paragraph(xml_escape(item), styles["body"]))
            elif "|" in item:
                story.append(Paragraph(xml_escape(item), styles["strong"]))
            else:
                story.append(Paragraph("- " + xml_escape(item), styles["bullet"]))

    doc.build(story)
    return path


if __name__ == "__main__":
    profile = load_profile()
    for payload in (ATS, STARTUP):
        print(build_resume(payload, profile))
        print(build_resume_pdf(payload, profile))
