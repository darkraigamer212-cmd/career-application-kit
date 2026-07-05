import json
import os
import re
from html import escape

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = os.path.abspath(os.path.dirname(__file__))
OUT = os.path.join(ROOT, "output", "tamil_report")
os.makedirs(OUT, exist_ok=True)

DATA = os.path.join(ROOT, "output", "pdf", "rathinam_rental_research_summary.json")
FONT_TA = os.path.join(ROOT, "assets", "NotoSansTamil-Regular.ttf")
FONT_TA_B = os.path.join(ROOT, "assets", "NotoSansTamil-Bold.ttf")
FONT_EN = os.path.join(ROOT, "assets", "NotoSans-Regular.ttf")
FONT_EN_B = os.path.join(ROOT, "assets", "NotoSans-Bold.ttf")

DATE_TA = "05 ஜூலை 2026"
NAVY = colors.HexColor("#123A63")
BLUE = colors.HexColor("#2E74B5")
LIGHT = colors.HexColor("#EAF2FB")
PALE = colors.HexColor("#F7FAFE")
GRID = colors.HexColor("#C7D3E3")
INK = colors.HexColor("#111827")
MUTED = colors.HexColor("#5B677A")


def read_houses():
    with open(DATA, "r", encoding="utf-8") as f:
        return json.load(f)["ranked"]


def money(v):
    if isinstance(v, int):
        return f"₹{v:,}"
    return str(v).replace("Rs ", "₹")


def register_fonts():
    pdfmetrics.registerFont(TTFont("Tamil", FONT_TA))
    pdfmetrics.registerFont(TTFont("TamilBold", FONT_TA_B))
    pdfmetrics.registerFont(TTFont("Latin", FONT_EN))
    pdfmetrics.registerFont(TTFont("LatinBold", FONT_EN_B))


def rich(text):
    text = escape(str(text))
    return re.sub(r"[A-Za-z0-9₹][A-Za-z0-9₹,.;:/()|&=\-+% ]*", lambda m: f'<font name="Latin">{m.group(0)}</font>', text)


def make_styles():
    register_fonts()
    base = getSampleStyleSheet()
    styles = {
        "cover": ParagraphStyle("cover", parent=base["Title"], fontName="TamilBold", fontSize=23, leading=32, alignment=TA_CENTER, textColor=NAVY),
        "meta": ParagraphStyle("meta", parent=base["BodyText"], fontName="Tamil", fontSize=10.5, leading=17, alignment=TA_CENTER, textColor=INK),
        "h1": ParagraphStyle("h1", parent=base["Heading1"], fontName="TamilBold", fontSize=15, leading=22, textColor=NAVY, spaceBefore=8, spaceAfter=7),
        "h2": ParagraphStyle("h2", parent=base["Heading2"], fontName="TamilBold", fontSize=12, leading=18, textColor=BLUE, spaceBefore=5, spaceAfter=5),
        "body": ParagraphStyle("body", parent=base["BodyText"], fontName="Tamil", fontSize=9.6, leading=15.5, textColor=INK, spaceAfter=5),
        "small": ParagraphStyle("small", parent=base["BodyText"], fontName="Tamil", fontSize=8.0, leading=12, textColor=INK),
        "tiny": ParagraphStyle("tiny", parent=base["BodyText"], fontName="Tamil", fontSize=6.7, leading=9, textColor=INK),
    }
    styles["smallW"] = ParagraphStyle("smallW", parent=styles["small"], textColor=colors.white)
    styles["tinyW"] = ParagraphStyle("tinyW", parent=styles["tiny"], textColor=colors.white)
    return styles


def P(text, style):
    return Paragraph(rich(text), style)


AREA_ROWS = [
    ["பகுதி", "பாதுகாப்பு", "தூரம்", "உணவு", "மருத்துவம்", "Internet", "பயணம்", "வாடகை", "அமைதி", "மொத்தம்"],
    ["Eachanari", "8", "10", "8", "8", "8", "10", "8", "8", "8.5"],
    ["Malumichampatti", "7", "8", "7", "7", "7", "8", "9", "8", "7.6"],
    ["Kurichi", "8", "7", "9", "8", "8", "7", "7", "7", "7.6"],
    ["Podanur", "7", "6", "9", "9", "8", "6", "7", "6", "7.2"],
    ["Kuniyamuthur", "8", "6", "9", "8", "9", "6", "6", "7", "7.4"],
]

MONTHLY = [
    ["வாடகை", 7000, 6500, 10000],
    ["மின்சாரம் (Electricity)", 1000, 600, 1500],
    ["இணையம் (Fiber Internet / Wi-Fi)", 900, 600, 1200],
    ["மளிகை + சமைப்பு", 5000, 3500, 6500],
    ["Bike fuel + maintenance", 1800, 1200, 2500],
    ["Emergency / medical buffer", 1500, 800, 2500],
]

ONE_TIME = [
    ["Security Deposit", "₹10,000", "Recommended house அடிப்படையில்"],
    ["Kitchen essentials", "₹3,500 - ₹6,000", "Induction/gas, பாத்திரங்கள், storage"],
    ["Mattress", "₹3,000 - ₹6,000", "வீட்டில் cot/bed இல்லையெனில்"],
    ["Wi-Fi setup/router", "₹0 - ₹2,000", "ISP offer பொறுத்து"],
    ["Basic utensils", "₹2,000 - ₹3,500", "Plate, tumbler, cooker, pan"],
    ["Cleaning supplies", "₹800 - ₹1,500", "Bucket, broom, mop, liquids"],
    ["Bike accessories", "₹1,000 - ₹2,000", "Lock, rain cover, basic toolkit"],
]

VERIFY = [
    "தண்ணீர் source - borewell / corporation / tanker backup",
    "Fiber Internet கிடைக்கிறதா - Airtel / Jio / ACT / local ISP",
    "அருகில் உள்ள neighbours எப்படி இருக்கிறார்கள்",
    "Bike parking பாதுகாப்பானதா",
    "Rent agreement, notice period, lock-in period",
    "Security Deposit refund terms எழுத்தில் உள்ளதா",
    "Power cuts எவ்வளவு வருகிறது",
    "இரவு street safety மற்றும் street light",
    "Noise - traffic, workshop, factory, loudspeaker போன்றவை",
    "Mobile signal indoor speed test",
    "Owner documents மற்றும் ID proof",
    "EB connection தனியாக உள்ளதா; meter reading எப்படி கணக்கிடப்படும்",
]

RISKS = [
    ["தனியாக வாழ்வது", "Daily routine, emergency contacts, location sharing வைத்துக்கொள்வது."],
    ["சமைப்பது", "Simple cooking plan வைத்து வெளி உணவை குறைப்பது."],
    ["Bills", "Rent, EB, Wi-Fi due dates-ஐ calendar reminder-ல் போடுவது."],
    ["Cleaning", "வாரம் இருமுறை room cleaning; மாதம் ஒருமுறை deep cleaning."],
    ["Maintenance", "Owner-க்கு WhatsApp proof உடன் repair record வைத்துக்கொள்வது."],
    ["Emergency", "Nearby hospital, pharmacy, நண்பர், owner contact numbers phone-ல் pin செய்வது."],
]


def table(rows, widths, styles, tiny=False):
    normal = styles["tiny"] if tiny else styles["small"]
    head = styles["tinyW"] if tiny else styles["smallW"]
    data = [[P(c, head if r == 0 else normal) for c in row] for r, row in enumerate(rows)]
    t = Table(data, colWidths=widths, repeatRows=1)
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), NAVY),
        ("GRID", (0, 0), (-1, -1), 0.35, GRID),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, PALE]),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    return t


def card(title, body_lines, styles):
    rows = [[P(title, styles["h2"])]] + [[P(x, styles["body"])] for x in body_lines]
    t = Table(rows, colWidths=[176 * mm])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), LIGHT),
        ("BOX", (0, 0), (-1, -1), 1, BLUE),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]))
    return t


def house_rows(houses):
    rows = [["#", "வீடு", "வாடகை", "Deposit", "தூரம்", "நேரம்", "Furnished", "Parking", "நன்மைகள்", "கவனிக்க வேண்டியது"]]
    for i, h in enumerate(houses, 1):
        kind = "1BHK House" if h["bedrooms"] == "1" and "House" in h["type"] else ("2BHK House" if h["bedrooms"] == "2" and "House" in h["type"] else f"{h['bedrooms']}BHK Flat")
        if i == 1:
            plus = "குறைந்த வாடகை, அருகில், Semi-Furnished"
            minus = "Water, parking, agreement verify செய்ய வேண்டும்"
        elif h["furnished"] == "Unfurnished":
            plus = "வாடகை கட்டுப்பாட்டில், commute acceptable"
            minus = "Unfurnished; setup செலவு வரும்"
        else:
            plus = "கல்லூரிக்கு அருகில், commute எளிது"
            minus = "Deposit / owner terms verify செய்ய வேண்டும்"
        if "500,000" in h["deposit"]:
            minus = "Deposit மிக அதிகம்; avoid/strict verify"
        rows.append([str(i), f"{h['locality']} - {kind}", money(h["rent"]), money(h["deposit"]), f"{h['distance']} km", h["time"], h["furnished"], h["parking"].replace("Ask owner", "கேட்க வேண்டும்"), plus, minus])
    return rows


def build_pdf(houses):
    s = make_styles()
    best = houses[0]
    monthly_now = sum(x[1] for x in MONTHLY)
    monthly_min = sum(x[2] for x in MONTHLY)
    monthly_max = sum(x[3] for x in MONTHLY)
    yearly = monthly_now * 12 + 17000
    path = os.path.join(OUT, "rathinam_house_report_tamil.pdf")
    doc = SimpleDocTemplate(path, pagesize=A4, leftMargin=16*mm, rightMargin=16*mm, topMargin=15*mm, bottomMargin=14*mm)
    story = []

    story += [Spacer(1, 30*mm), P("ரத்தினம் கல்லூரி அருகிலுள்ள வாடகை வீடு", s["cover"]), P("முழுமையான ஆய்வு அறிக்கை", s["cover"]), Spacer(1, 12*mm)]
    story += [P("Prepared for: My Father", s["meta"]), P("Prepared by: Dragon", s["meta"]), P(f"Date: {DATE_TA}", s["meta"]), Spacer(1, 18*mm)]
    story += [card("முக்கிய முடிவு", [
        "சுமார் ஒரு வருடத்திற்கு மட்டும், குறைந்த செலவு + அமைதியான சூழல் + விரைவான commute என்ற அடிப்படையில் Eachanari பகுதியில் உள்ள 1 BHK Residential House சிறந்த தேர்வாக தெரிகிறது.",
        f"வாடகை: {money(best['rent'])} | Security Deposit: {money(best['deposit'])} | பயணம்: {best['distance']} km / {best['time']}",
    ], s), Spacer(1, 5*mm), P("இந்த அறிக்கை உணர்ச்சிவசப்பட்ட முடிவு அல்ல. Public rental portals-ல் கிடைத்த தகவல்கள், budget, commute, பாதுகாப்பு, internet, future MBA plan ஆகியவற்றை வைத்து practical ஆய்வாக தயாரிக்கப்பட்டது.", s["body"]), PageBreak()]

    story += [P("1. Hostel-லிருந்து வெளியேற விரும்பும் காரணம்", s["h1"])]
    story.append(P("இந்த மாற்றத்தின் நோக்கம் வெறும் வசதி அல்ல. அமைதியான சூழலில் படிப்பு, AI development, software projects, MBA preparation ஆகியவற்றில் கவனம் செலுத்த சுமார் 10-12 மாதங்களுக்கு மட்டும் independent room தேவைப்படுகிறது.", s["body"]))
    for item in ["Better concentration", "Quiet environment", "Better sleep", "Better food control", "AI development மற்றும் software projects", "MBA preparation", "Privacy மற்றும் independence", "இது ஒரு வருட temporary plan மட்டும்"]:
        story.append(P("- " + item, s["body"]))
    story += [P("2. என் தேவைகள்", s["h1"])]
    for item in ["₹5,000-₹10,000 rent; மிக நல்ல option என்றால் மட்டும் ₹12,000 வரை", "Independent room / 1 RK / 1 BHK / small independent house", "No PG, No Hostel, No shared room", "Safe area, peaceful neighbourhood, good roads", "Hospital, grocery, pharmacy அருகில்", "Bike travel 15-20 minutes-க்கு கீழ்", "Good internet, good water, good mobile signal", "ஒரு வருடத்திற்கு பிறகு easy move-out"]:
        story.append(P("- " + item, s["body"]))
    story += [P("3. பகுதிகள் ஒப்பீடு", s["h1"]), table(AREA_ROWS, [30*mm,16*mm,14*mm,14*mm,18*mm,17*mm,16*mm,16*mm,17*mm,18*mm], s, tiny=True), Spacer(1,4*mm), P("மதிப்பீடு public information, commute logic, residential suitability, facilities availability ஆகியவற்றின் அடிப்படையில் 10-ல் கணிக்கப்பட்டது. நேரில் பார்த்த பிறகு final confirmation அவசியம்.", s["small"]), PageBreak()]

    story += [P("4. Top 10 வீடுகள்", s["h1"]), table(house_rows(houses), [7*mm,33*mm,15*mm,18*mm,13*mm,14*mm,18*mm,19*mm,30*mm,29*mm], s, tiny=True), PageBreak()]
    story += [P("5. பரிந்துரைக்கப்படும் வீடு", s["h1"]), card("Recommended House - Eachanari 1 BHK Residential House", [
        f"Rent: {money(best['rent'])}",
        f"Security Deposit: {money(best['deposit'])}",
        f"Bike distance: {best['distance']} km",
        f"Travel time: {best['time']}",
        f"Furnishing: {best['furnished']}",
        "கல்லூரிக்கு அருகில் இருப்பதால் daily commute சுலபம்.",
        "Semi-Furnished என்பதால் ஆரம்ப setup செலவு குறையும்.",
        "Temporary one-year stay-க்கு rent மற்றும் deposit இரண்டும் practical ஆக இருக்கிறது.",
    ], s), P("இந்த வீடு அதிக வசதிக்காக அல்ல; குறைந்த செலவில் அமைதியான, productive சூழலை உருவாக்குவதற்காக தேர்வு செய்யப்படுகிறது. Final advance கொடுக்கும் முன் water, parking, agreement, deposit refund, neighbours, internet availability ஆகியவை நேரில் verify செய்ய வேண்டும்.", s["body"])]

    budget = [["செலவு", "தற்போதைய", "Minimum", "Maximum"]] + [[x[0], money(x[1]), money(x[2]), money(x[3])] for x in MONTHLY] + [["மொத்தம்", money(monthly_now), money(monthly_min), money(monthly_max)]]
    story += [P("6. Expected Expenses - Monthly Budget", s["h1"]), table(budget, [62*mm,38*mm,38*mm,38*mm], s), P(f"Yearly cost: மாத சராசரி {money(monthly_now)} × 12 + basic setup buffer சேர்த்து சுமார் {money(yearly)}.", s["body"]), PageBreak()]
    story += [P("7. One-Time Expenses", s["h1"]), table([["Item", "Approx. Amount", "குறிப்பு"]] + ONE_TIME + [["மொத்தம்", "₹21,300 - ₹31,000", "Deposit உட்பட approximate setup range"]], [55*mm,42*mm,79*mm], s), PageBreak()]

    story += [P("8. Advance கொடுக்கும் முன் verify செய்ய வேண்டியவை", s["h1"])]
    for item in VERIFY:
        story.append(P("- " + item, s["body"]))
    story += [P("9. Advantages", s["h1"]), P("Independent room எனக்கு hostel-ஐ விட better focus, privacy, health, projects, MBA preparation, time saving ஆகியவற்றுக்கு உதவும். Hostel-ல் கிடைக்கும் social environment நல்லதே; ஆனால் இந்த ஒரு வருடத்தில் எனக்கு தேவையானது அமைதியான execution environment.", s["body"])]
    story += [P("10. Possible Risks", s["h1"]), table([["Risk", "Manage செய்யும் முறை"]] + RISKS, [52*mm,124*mm], s)]
    story += [P("11. Final Recommendation", s["h1"]), card("அப்பாவுக்கு மரியாதையான வேண்டுகோள்", ["அப்பா, இந்த முடிவை நான் உணர்ச்சிவசப்பட்டு எடுக்கவில்லை. கல்லூரிக்கு தூரம், வாடகை, deposit, பாதுகாப்பு, internet, தண்ணீர், future MBA plan ஆகியவற்றை வைத்து கவனமாக ஆய்வு செய்தேன். இந்த வீடு சுமார் ஒரு வருடத்திற்கு மட்டும். Graduation முடிந்த பிறகு MBA abroad நோக்கி நகரும் திட்டம் உள்ளது. அதனால் இது நிரந்தர செலவு அல்ல; better productivity, health, projects, preparation ஆகியவற்றுக்கான temporary investment. நீங்கள் பார்த்து உங்கள் கருத்தை சொல்லுங்கள். நீங்கள் சரி என்று நினைத்த பிறகே advance கொடுப்போம்."], s)]

    def footer(canvas, doc_):
        canvas.saveState()
        canvas.setFont("Tamil", 7)
        canvas.setFillColor(MUTED)
        canvas.drawString(16*mm, 8*mm, "ரத்தினம் கல்லூரி அருகிலுள்ள வாடகை வீடு - ஆய்வு அறிக்கை")
        canvas.drawRightString(194*mm, 8*mm, str(doc_.page))
        canvas.restoreState()

    doc.build(story, onFirstPage=footer, onLaterPages=footer)
    return path


def doc_run(p, text, bold=False, size=10.5, color="111827"):
    r = p.add_run(text)
    r.font.name = "Noto Sans Tamil"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans Tamil")
    r._element.rPr.rFonts.set(qn("w:cs"), "Noto Sans Tamil")
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = RGBColor.from_string(color)


def doc_p(doc, text, bold=False, size=10.5, color="111827", center=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(5)
    doc_run(p, text, bold, size, color)
    return p


def shade(cell, fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def doc_table(doc, rows):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = t.cell(i, j)
            c.text = ""
            p = c.paragraphs[0]
            doc_run(p, str(val), bold=(i == 0), size=8, color=("FFFFFF" if i == 0 else "111827"))
            if i == 0:
                shade(c, "123A63")
            elif i % 2 == 0:
                shade(c, "F7FAFE")
    doc.add_paragraph()


def build_docx(houses):
    best = houses[0]
    monthly_now = sum(x[1] for x in MONTHLY)
    yearly = monthly_now * 12 + 17000
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(1.7); sec.bottom_margin = Cm(1.5); sec.left_margin = Cm(1.7); sec.right_margin = Cm(1.7)
    doc_p(doc, "ரத்தினம் கல்லூரி அருகிலுள்ள வாடகை வீடு", True, 22, "123A63", True)
    doc_p(doc, "முழுமையான ஆய்வு அறிக்கை", True, 18, "123A63", True)
    doc_p(doc, f"Prepared for: My Father\nPrepared by: Dragon\nDate: {DATE_TA}", False, 11, "111827", True)
    doc.add_page_break()
    for title, body in [
        ("1. Hostel-லிருந்து வெளியேற விரும்பும் காரணம்", "Better concentration, quiet environment, better sleep, better food control, AI development, software projects, MBA preparation, privacy, independence ஆகியவற்றுக்காக சுமார் ஒரு வருடத்திற்கு independent room தேவைப்படுகிறது."),
        ("2. என் தேவைகள்", "₹5,000-₹10,000 rent, Independent room / 1 RK / 1 BHK, No PG, No Hostel, safe area, peaceful neighbourhood, hospital/grocery அருகில், bike travel 15-20 minutes-க்கு கீழ், நல்ல internet/water, easy move-out."),
    ]:
        doc_p(doc, title, True, 15, "123A63"); doc_p(doc, body)
    doc_p(doc, "3. பகுதிகள் ஒப்பீடு", True, 15, "123A63"); doc_table(doc, AREA_ROWS)
    doc_p(doc, "4. Top 10 வீடுகள்", True, 15, "123A63"); doc_table(doc, house_rows(houses))
    doc_p(doc, "5. பரிந்துரைக்கப்படும் வீடு", True, 15, "123A63")
    doc_p(doc, f"{best['title']} - Rent {money(best['rent'])}, Deposit {money(best['deposit'])}, {best['distance']} km / {best['time']}, {best['furnished']}. இது temporary stay-க்கு best balance.")
    budget = [["செலவு", "தற்போதைய", "Minimum", "Maximum"]] + [[x[0], money(x[1]), money(x[2]), money(x[3])] for x in MONTHLY]
    doc_p(doc, "6. Expected Expenses", True, 15, "123A63"); doc_table(doc, budget); doc_p(doc, f"Yearly cost: சுமார் {money(yearly)}.")
    doc_p(doc, "7. One-Time Expenses", True, 15, "123A63"); doc_table(doc, [["Item", "Approx. Amount", "குறிப்பு"]] + ONE_TIME)
    doc_p(doc, "8. Verify Before Advance", True, 15, "123A63")
    for item in VERIFY: doc_p(doc, "- " + item)
    doc_p(doc, "9. Advantages", True, 15, "123A63"); doc_p(doc, "Independent room hostel-ஐ விட better focus, privacy, health, projects, MBA preparation, time saving ஆகியவற்றுக்கு உதவும்.")
    doc_p(doc, "10. Possible Risks", True, 15, "123A63"); doc_table(doc, [["Risk", "Manage செய்யும் முறை"]] + RISKS)
    doc_p(doc, "11. Final Recommendation", True, 15, "123A63"); doc_p(doc, "அப்பா, இந்த முடிவை நான் உணர்ச்சிவசப்பட்டு எடுக்கவில்லை. இது சுமார் ஒரு வருடத்திற்கு மட்டும்; graduation-க்குப் பிறகு MBA abroad plan உள்ளது. உங்கள் கருத்தை மரியாதையுடன் கேட்கிறேன்.")
    path = os.path.join(OUT, "rathinam_house_report_tamil.docx")
    doc.save(path)
    return path


def build_md_and_whatsapp(houses):
    best = houses[0]
    monthly_now = sum(x[1] for x in MONTHLY)
    yearly = monthly_now * 12 + 17000
    md = [
        "# ரத்தினம் கல்லூரி அருகிலுள்ள வாடகை வீடு",
        "## முழுமையான ஆய்வு அறிக்கை",
        "",
        "Prepared for: My Father  ",
        "Prepared by: Dragon  ",
        f"Date: {DATE_TA}",
        "",
        "## முக்கிய முடிவு",
        f"Recommended: {best['title']} - Rent {money(best['rent'])}, Deposit {money(best['deposit'])}, {best['distance']} km / {best['time']}.",
        "",
        "## Sections",
        "1. Hostel-லிருந்து வெளியேற காரணம் - focus, sleep, food control, AI projects, MBA preparation.",
        "2. Requirements - No PG/Hostel, independent room, safe area, good internet/water.",
        "3. Area comparison - Eachanari highest overall for commute.",
        "4. Top 10 houses - see table data in PDF/DOCX.",
        "5. Recommended house - Eachanari 1BHK Residential House.",
        f"6. Monthly budget - expected {money(monthly_now)}; yearly estimate {money(yearly)}.",
        "7. One-time expenses - deposit, kitchen, mattress, Wi-Fi, utensils, cleaning, bike accessories.",
        "8. Verify before advance - water, internet, neighbours, parking, agreement, deposit refund, power cuts, safety, EB.",
        "9. Advantages - focus, privacy, health, projects, MBA preparation, time saving.",
        "10. Risks - living alone, cooking, bills, cleaning, maintenance, emergency; manageable with routine.",
        "11. Final recommendation - respectful request to father for opinion.",
    ]
    md_path = os.path.join(OUT, "rathinam_house_report_tamil.md")
    with open(md_path, "w", encoding="utf-8") as f:
        f.write("\n".join(md))
    wa = f"""ரத்தினம் கல்லூரி அருகில் வாடகை வீடு - சுருக்கம்

அப்பா, இது சுமார் 10-12 மாதங்களுக்கு மட்டும் temporary stay.

Recommended:
{best['title']}
Location: {best['address']}
Rent: {money(best['rent'])}
Security Deposit: {money(best['deposit'])}
Distance: {best['distance']} km
Bike time: {best['time']}
Furnishing: {best['furnished']}

ஏன் இது நல்லது:
- கல்லூரிக்கு மிக அருகில்
- குறைந்த rent + குறைந்த deposit
- Semi-Furnished; setup செலவு குறையும்
- AI projects, software work, MBA preparation-க்கு அமைதியான சூழல்
- Hostel-ஐ விட privacy, focus, sleep, food control மேம்படும்

Monthly expected cost: {money(monthly_now)}
One-year estimate: {money(yearly)}

Advance கொடுக்கும் முன் verify:
Water, Fiber Internet, parking, neighbours, night safety, rent agreement, deposit refund, power cuts, EB meter, mobile signal.

இது emotional decision இல்லை. Carefully researched temporary productivity investment. நீங்கள் பார்த்து உங்கள் கருத்து சொல்லுங்கள்."""
    wa_path = os.path.join(OUT, "whatsapp_one_page_summary_tamil.txt")
    with open(wa_path, "w", encoding="utf-8") as f:
        f.write(wa)
    return md_path, wa_path


if __name__ == "__main__":
    houses = read_houses()
    print("pdf:", build_pdf(houses))
    print("docx:", build_docx(houses))
    md, wa = build_md_and_whatsapp(houses)
    print("markdown:", md)
    print("whatsapp:", wa)
